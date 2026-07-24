from docx import Document
from docx.shared import Pt

def create_manuscript_outline():
    doc = Document()

    # Title
    title = doc.add_heading('Manuscript Outline and FAIR Data Management Plan', 0)

    # Part 1: Manuscript Outline
    doc.add_heading('Part 1: Manuscript Outline', level=1)

    doc.add_heading('Preliminary Title', level=2)
    doc.add_paragraph('Multi-Omics Integration of Microarray and RNAseq Data to Characterize Plant Responses to Space Radiation Stress')

    doc.add_heading('Abstract', level=2)
    doc.add_paragraph('As humanity extends its presence in space, understanding how plants adapt to ionizing radiation becomes critical. This study integrates legacy microarray data with newly available RNAseq data from OSDR to identify robust molecular signatures of radiation response in Arabidopsis and Rice. We employ deep learning methods to predict stress-responsive pathways and discuss the implications for crop improvement in space habitats.')

    doc.add_heading('Introduction', level=2)
    doc.add_paragraph('- Importance of plants in long-duration space missions.')
    doc.add_paragraph('- Challenges posed by ionizing radiation (HZE, GCR).')
    doc.add_paragraph('- Transition from microarray to RNAseq in space biology.')
    doc.add_paragraph('- Objectives: Identify conserved radiation-responsive genes across datasets.')

    doc.add_heading('Methods', level=2)
    doc.add_paragraph('- Data Retrieval: OSDR and GeneLab accessions (GLDS-7, 37, 38, 46, 120, 329, 321, 320, 296, 282, 281, 223).')
    doc.add_paragraph('- Preprocessing: Normalization and quality control for microarray and RNAseq data.')
    doc.add_paragraph('- Differential Expression Analysis (DEG): DESeq2 for RNAseq, Limma for microarray.')
    doc.add_paragraph('- Network Analysis: WGCNA for module identification.')
    doc.add_paragraph('- Deep Learning: Convolutional Neural Networks (CNNs) or Transformers for gene expression prediction (potential).')

    doc.add_heading('Results', level=2)
    doc.add_paragraph('- Overlap between microarray and RNAseq signatures.')
    doc.add_paragraph('- Functional enrichment (KEGG, SBGNview) of common pathways.')
    doc.add_paragraph('- Species-specific responses (Arabidopsis vs. Rice).')

    doc.add_heading('Discussion', level=2)
    doc.add_paragraph('- Mechanisms of DNA repair and ROS scavenging.')
    doc.add_paragraph('- Translation of findings to crop resilience.')
    doc.add_paragraph('- Limitations and future directions.')

    # Part 2: FAIR Data Management Plan
    doc.add_page_break()
    doc.add_heading('Part 2: FAIR Data Management Plan', level=1)

    doc.add_heading('1. Findable', level=2)
    doc.add_paragraph('- Data will be deposited in the Zenodo repository with a unique Digital Object Identifier (DOI).')
    doc.add_paragraph('- Metadata will follow the OSDR (Open Science Data Repository) standards.')
    doc.add_paragraph('- Keywords: Space Biology, Plant Radiation, Transcriptomics, FAIR.')

    doc.add_heading('2. Accessible', level=2)
    doc.add_paragraph('- Data and code will be publicly available under a Creative Commons Attribution (CC BY) license.')
    doc.add_paragraph('- The GitHub repository will be linked to Zenodo for automatic versioning.')

    doc.add_heading('3. Interoperable', level=2)
    doc.add_paragraph('- Use of standard file formats: CSV for tabular data, XLSX for combined results, Rmd/R scripts for analysis.')
    doc.add_paragraph('- Gene identifiers will be mapped to universal databases (NCBI, TAIR).')

    doc.add_heading('4. Reusable', level=2)
    doc.add_paragraph('- Comprehensive documentation (README) explaining the repository structure and script usage.')
    doc.add_paragraph('- Inclusion of raw and processed data to allow for re-analysis.')
    doc.add_paragraph('- Clear licensing (LICENSE file) in the repository.')

    # Save the document
    doc.save('docs/Manuscript_Outline_and_FAIR_Plan.docx')
    print("Document 'docs/Manuscript_Outline_and_FAIR_Plan.docx' created successfully.")

if __name__ == '__main__':
    create_manuscript_outline()
